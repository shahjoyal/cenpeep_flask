"""
generate_model_report.py — CENPEEP field-detection decision-trace report
============================================================================
One-off script: runs the SAME parsing/detection pipeline used by
routes/upload.py against a real Excel file, with full decision tracing
turned on (which stage matched each column, why, with what confidence,
against which training phrase), and writes a Word document containing
ONLY, for each sheet: the sheet name as a heading, followed by its
column-by-column decision table. No title page, no methodology sections,
no summary counts, no legend.

USAGE
-----
    python generate_model_report.py path/to/your_workbook.xlsx

Produces "Model_Decision_Report.docx" in the same folder as this script.

This is a reporting/documentation tool, not part of the running app —
nothing here is imported by app.py or routes/upload.py. Once you've
generated the report you want, you can leave this file in the repo (it
does nothing unless run directly) or just stop running it. The bottom of
the file has a single guarded entry point:

    if __name__ == "__main__":
        main()

Comment out the body of main() (or the call inside it) if you want to keep
the file around purely for reference without ever re-running it.
"""

import sys
import os
import io
import statistics

import openpyxl

# Reuse the REAL production internals so this report reflects exactly what
# the running app does — not a re-description of it.
from routes.upload import (
    _is_readable_worksheet,
    _find_header_row,
    _label_to_field,
    _sym_to_field,
    _to_num,
    _parse_cenpeep_layout,
    HEADER_SCAN_ROWS,
)
from ml.training_data import is_non_field_header
from ml.field_classifier import get_classifier, DEFAULT_CONFIDENCE_THRESHOLD


# ── Step 1: run the pipeline with full tracing ──────────────────────────────
def trace_workbook(path):
    """
    Same job as routes.upload.parse_workbook(), but every column's decision
    is recorded (not just the final extracted values), for reporting.
    """
    filename = os.path.basename(path)
    with open(path, "rb") as f:
        file_bytes = f.read()

    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
    clf = get_classifier()

    sheet_traces = []
    for name in wb.sheetnames:
        ws = wb[name]
        if not _is_readable_worksheet(ws):
            sheet_traces.append({
                "sheetName": name,
                "skipped": True,
                "reason": "Not a data worksheet (e.g. a chart sheet) — no rows to read.",
                "columns": [],
            })
            continue

        rows = [list(r) for r in ws.iter_rows(values_only=True)]

        # Strategy 1: strict CenPeep column layout — check first, same
        # priority as production.
        ext1, raw1 = _parse_cenpeep_layout(rows)
        if len(ext1) >= 5:
            sheet_traces.append({
                "sheetName": name,
                "skipped": False,
                "strategy": "cenpeep_column",
                "columns": [
                    {
                        "header": r["particulars"] or r["symbol"],
                        "decision": "matched",
                        "method": "cenpeep_column_layout",
                        "fieldId": _sym_to_field(r["symbol"]),
                        "confidence": 1.0,
                        "matchedPhrase": None,
                        "readings": 1,
                        "average": r["value"],
                    }
                    for r in raw1
                ],
            })
            continue

        # Strategy 2: raw tabular layout — find header row, trace every column.
        sample = rows[:HEADER_SCAN_ROWS]
        header_row_idx = _find_header_row(sample)
        if header_row_idx is None:
            sheet_traces.append({
                "sheetName": name, "skipped": True,
                "reason": "No header row could be identified in the first "
                          f"{HEADER_SCAN_ROWS} rows.",
                "columns": [],
            })
            continue

        headers = rows[header_row_idx]
        data_rows = rows[header_row_idx + 1:]

        col_traces = []
        unmatched_idx, unmatched_text = [], []

        for col_idx, hdr in enumerate(headers):
            if hdr is None or not str(hdr).strip():
                continue
            if is_non_field_header(hdr):
                col_traces.append({
                    "header": str(hdr), "decision": "excluded",
                    "method": "structural_header_list",
                    "fieldId": None, "confidence": None, "matchedPhrase": None,
                })
                continue

            fid = _label_to_field(str(hdr))
            if fid:
                col_traces.append({
                    "header": str(hdr), "decision": "matched",
                    "method": "rule_based_alias", "fieldId": fid,
                    "confidence": 1.0, "matchedPhrase": None,
                    "colIdx": col_idx,
                })
            else:
                unmatched_idx.append(col_idx)
                unmatched_text.append(str(hdr))
                col_traces.append({
                    "header": str(hdr), "decision": "pending_ml",
                    "method": None, "fieldId": None, "confidence": None,
                    "matchedPhrase": None, "colIdx": col_idx,
                })

        if unmatched_text:
            preds = clf.predict_batch(unmatched_text, threshold=DEFAULT_CONFIDENCE_THRESHOLD)
            pred_by_text = dict(zip(unmatched_text, preds))
            for ct in col_traces:
                if ct["decision"] != "pending_ml":
                    continue
                fid, score, matched_example = pred_by_text[ct["header"]]
                if fid:
                    ct.update(decision="matched", method="ml_tfidf_cosine",
                              fieldId=fid, confidence=round(score, 3),
                              matchedPhrase=matched_example)
                elif matched_example is not None:
                    ct.update(decision="rejected_out_of_scope", method="ml_tfidf_cosine",
                              confidence=round(score, 3), matchedPhrase=matched_example)
                else:
                    ct.update(decision="rejected_low_confidence", method="ml_tfidf_cosine",
                              confidence=round(score, 3))

        # Accumulate values + averages per matched field, same as production.
        field_values = {}
        for ct in col_traces:
            if ct["decision"] != "matched":
                continue
            fid, col_idx = ct["fieldId"], ct["colIdx"]
            vals = []
            for row in data_rows:
                v = row[col_idx] if col_idx < len(row) else None
                num = _to_num(v)
                if num is not None:
                    vals.append(num)
            ct["readings"] = len(vals)
            ct["average"] = round(statistics.mean(vals), 4) if vals else None
            if not vals:
                ct["decision"] = "matched_but_no_numeric_data"

        sheet_traces.append({
            "sheetName": name, "skipped": False,
            "strategy": "raw_tabular_ml", "columns": col_traces,
        })

    wb.close()
    return {"filename": filename, "sheets": sheet_traces}


# ── Step 2: build the Word document ─────────────────────────────────────────
def build_report(trace, output_path):
    from docx import Document
    from docx.shared import Pt
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

    doc = Document()

    # -- base style
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    def shade_cell(cell, hex_color):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = tcPr.makeelement(qn("w:shd"), {
            qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): hex_color,
        })
        tcPr.append(shd)

    decision_colors = {
        "matched": "D9EAD3",
        "matched_but_no_numeric_data": "FFF2CC",
        "excluded": "F3F3F3",
        "rejected_out_of_scope": "F4CCCC",
        "rejected_low_confidence": "FCE5CD",
    }

    for sh in trace["sheets"]:
        doc.add_heading(sh["sheetName"], level=2)
        if sh.get("skipped"):
            doc.add_paragraph(f"Skipped — {sh['reason']}")
            continue
        if not sh["columns"]:
            doc.add_paragraph("No fields detected in this sheet.")
            continue

        table = doc.add_table(rows=1, cols=7)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Light Grid Accent 1"
        hdr_cells = table.rows[0].cells
        for i, txt in enumerate(["Column Header", "Decision", "Method", "Field", "Confidence",
                                  "Matched Against", "Avg (readings)"]):
            hdr_cells[i].text = txt
            hdr_cells[i].paragraphs[0].runs[0].bold = True

        for c in sh["columns"]:
            row = table.add_row().cells
            row[0].text = str(c["header"])[:60]
            row[1].text = c["decision"].replace("_", " ")
            row[2].text = (c.get("method") or "—").replace("_", " ")
            row[3].text = c.get("fieldId") or "—"
            row[4].text = f"{c['confidence']:.2f}" if c.get("confidence") is not None else "—"
            # For rule-based matches there's no "closest training phrase" —
            # the header matched a hand-built alias/symbol exactly, so show
            # that instead of leaving it blank. For ML matches, this is the
            # actual training-set example whose similarity score won/lost
            # the confidence gate — i.e. exactly what the header was
            # compared against to reach this decision.
            if c.get("method") == "rule_based_alias":
                row[5].text = "exact alias/symbol lookup"
            elif c.get("matchedPhrase"):
                row[5].text = str(c["matchedPhrase"])[:60]
            else:
                row[5].text = "—"
            if c.get("readings") is not None and c.get("average") is not None:
                row[6].text = f"{c['average']:.3g}  (n={c['readings']})"
            else:
                row[6].text = "—"
            fill = decision_colors.get(c["decision"], "FFFFFF")
            for cell in row:
                shade_cell(cell, fill)
        doc.add_paragraph()

    doc.save(output_path)


def main():
    if len(sys.argv) < 2:
        print("Usage: python generate_model_report.py path/to/workbook.xlsx")
        sys.exit(1)

    excel_path = sys.argv[1]
    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Model_Decision_Report.docx")

    print(f"Tracing decisions for: {excel_path}")
    trace = trace_workbook(excel_path)

    print(f"Building report: {output_path}")
    build_report(trace, output_path)

    print("Done.")


if __name__ == "__main__":
    main()