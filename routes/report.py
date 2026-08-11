"""
report.py — on-demand "Field Detection Report" (browser-triggered)
====================================================================
This is the r.py (generate_model_report.py) table — Field / Detected
From / Method / Confidence / Value, plus the "Fields Not Detected"
list — wired up as a real API route instead of a CLI-only script, so
the "Results & Losses" tab can offer a download button for it.

It does NOT re-parse the workbook. The frontend already has everything
a report needs from the /api/upload response (fieldDetail, extracted,
missingFields, primarySheet) — this route just renders that straight
into a .docx, exactly like build_report() in r.py does.

DATE-WISE PROCESSES
--------------------
Field *detection* (which header a field came from, rule vs ML, the
confidence) doesn't change when the file is sliced by date — it's the
same sheet, same columns. Only the *value* each field resolves to
changes, because it's now averaged over just the rows in that
process's date range instead of the whole file. So when the request
includes a `processes` list, this builds one section per process —
same field-detail table, but with that process's averaged values
substituted in (falling back to the whole-file value for any field
the dated log doesn't carry, e.g. manual-only fields) — instead of
re-deriving detection per process, which wouldn't mean anything.

No processes in the request → one section for the whole file,
identical in shape to what r.py has always produced.
"""

import io
import re
from flask import Blueprint, request, send_file, jsonify

report_bp = Blueprint('report', __name__)

# Same palette/labels as r.py's build_report(), kept in sync on purpose —
# this route is meant to be the same report, just reachable from the UI.
METHOD_COLORS = {
    "cenpeep_column": "D9EAD3",
    "rule": "D9EAD3",
    "ml": "D6E4F0",
    "derived_fallback": "FCE5CD",
}
METHOD_LABELS = {
    "cenpeep_column": "CenPeep layout (exact)",
    "rule": "Exact alias/symbol match",
    "ml": "AI (ML) match",
    "derived_fallback": "Defaulted from another field",
}


def _shade_cell(cell, hex_color):
    from docx.oxml.ns import qn
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.makeelement(qn("w:shd"), {
        qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): hex_color,
    })
    tcPr.append(shd)


def _add_field_section(doc, heading, field_detail, extracted, missing_fields, level=2):
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc.add_heading(heading, level=level)

    if not field_detail:
        doc.add_paragraph("No fields could be detected on the selected sheet.")
    else:
        table = doc.add_table(rows=1, cols=5)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Light Grid Accent 1"
        hdr_cells = table.rows[0].cells
        for i, txt in enumerate(["Field", "Detected From", "Method", "Confidence", "Value"]):
            hdr_cells[i].text = txt
            hdr_cells[i].paragraphs[0].runs[0].bold = True

        for fid in sorted(field_detail.keys(), key=lambda f: field_detail[f].get("label", f) or f):
            d = field_detail[fid] or {}
            row = table.add_row().cells
            row[0].text = d.get("label") or fid
            row[1].text = d.get("header") or "—"
            method = d.get("source", "rule")
            row[2].text = METHOD_LABELS.get(method, method)
            conf = d.get("confidence")
            row[3].text = f"{conf:.2f}" if isinstance(conf, (int, float)) else "—"
            val = extracted.get(fid)
            row[4].text = f"{val:.4g}" if isinstance(val, (int, float)) else str(val or "—")
            fill = METHOD_COLORS.get(method, "FFFFFF")
            for cell in row:
                _shade_cell(cell, fill)

    doc.add_paragraph()
    doc.add_heading("Fields Not Detected", level=level + 1)
    if missing_fields:
        doc.add_paragraph(
            "The following required CENPEEP input fields were not found on "
            "any sheet in this workbook and must be entered manually:"
        )
        for m in missing_fields:
            label = (m.get("label") or m.get("id")) if isinstance(m, dict) else str(m)
            doc.add_paragraph(label or str(m), style="List Bullet")
    else:
        doc.add_paragraph("All required CENPEEP input fields were detected.")
    doc.add_paragraph()


@report_bp.route('', methods=['POST'])
def generate_report():
    """
    Body (JSON):
      {
        "filename": str,
        "primarySheet": str,
        "fieldDetail":   {fieldId: {label, header, source, confidence}},
        "extracted":     {fieldId: value},          // whole-file values
        "missingFields": [{id, label}, ...],
        "processes": [                               // optional
          {"title": str, "start": str|None, "end": str|None,
           "rowCount": int, "avg": {fieldId: value}},
          ...
        ]
      }
    Returns the .docx as a file download.
    """
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        return jsonify({
            'ok': False,
            'error': 'python-docx is not installed on the server (add it to requirements.txt).',
        }), 500

    data = request.get_json(force=True, silent=True) or {}
    field_detail   = data.get('fieldDetail') or {}
    extracted      = data.get('extracted') or {}
    missing_fields = data.get('missingFields') or []
    primary_sheet  = data.get('primarySheet') or ''
    filename       = data.get('filename') or 'workbook'
    processes      = data.get('processes') or []

    if not field_detail:
        return jsonify({
            'ok': False,
            'error': 'No field detail to report on — upload and parse a file first.',
        }), 400

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    doc.add_heading(f"CENPEEP Field Detection Report — {filename}", level=1)

    if not processes:
        # Whole-file report — identical shape to r.py's output.
        _add_field_section(doc, primary_sheet or "No sheet selected",
                            field_detail, extracted, missing_fields)
    else:
        # One section per date-wise process, each using its own averaged
        # values but the same field-detection info (see module docstring).
        for p in processes:
            title      = p.get('title') or 'Process'
            start      = p.get('start') or '—'
            end        = p.get('end') or '—'
            row_count  = p.get('rowCount', 0)
            avg        = p.get('avg') or {}
            proc_extracted = {**extracted, **avg}
            heading = f"{title}  —  {start} \u2192 {end}  ({row_count} row{'s' if row_count != 1 else ''})"
            _add_field_section(doc, heading, field_detail, proc_extracted, missing_fields)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    base = filename.rsplit('.', 1)[0] if '.' in filename else filename
    safe_name = re.sub(r'[^A-Za-z0-9 _-]', '', base).strip() or 'workbook'
    out_name = f"CENPEEP_Field_Report_{safe_name}.docx"

    return send_file(
        buf,
        as_attachment=True,
        download_name=out_name,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
